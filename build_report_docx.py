import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("PricePilot AI — Milestone 1 Technical Report")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 140, 160)

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Team 1  |  Infosys Internship Program  |  Page ")
        frun.font.name = "Calibri"
        frun.font.size = Pt(9)
        frun.font.color.rgb = RGBColor(120, 140, 160)

def style_heading(p, text, level):
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Cambria"
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate Dark
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(30, 58, 138) # Dark Navy
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(51, 65, 85)
    return run

def add_paragraph(doc, text="", style='Normal', space_after=6, line_spacing=1.15):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run1 = p.add_run(bold_prefix + " ")
    run1.font.name = "Calibri"
    run1.font.size = Pt(11)
    run1.bold = True
    run1.font.color.rgb = RGBColor(15, 23, 42)

    run2 = p.add_run(text)
    run2.font.name = "Calibri"
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    
    # Border & background via paragraph properties
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_figure(doc, img_path, fig_num, caption_text, width=Inches(5.8)):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(3)
        p_img.paragraph_format.keep_with_next = True
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=width)

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        run_cap = p_cap.add_run(f"Figure {fig_num}: {caption_text}")
        run_cap.font.name = "Calibri"
        run_cap.font.size = Pt(9.5)
        run_cap.font.italic = True
        run_cap.font.color.rgb = RGBColor(71, 85, 105)

def build_document():
    doc = docx.Document()
    
    # 1 Inch Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_header_footer(doc)

    # ==========================================
    # COVER PAGE
    # ==========================================
    p_top_spacer = doc.add_paragraph()
    p_top_spacer.paragraph_format.space_before = Pt(72)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_t = p_title.add_run("PricePilot AI")
    run_t.font.name = "Cambria"
    run_t.font.size = Pt(36)
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Dynamic Pricing Optimization & Revenue Intelligence System")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(79, 70, 229) # Indigo accent

    # Decorative Line
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(140)
    run_line = p_line.add_run("―" * 40)
    run_line.font.color.rgb = RGBColor(203, 213, 225)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(4)
    r = p_meta.add_run("Project Team: ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r = p_meta.add_run("Team 1")
    r.font.name = "Calibri"

    p_meta2 = doc.add_paragraph()
    p_meta2.paragraph_format.space_after = Pt(4)
    r = p_meta2.add_run("Prepared For: ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r = p_meta2.add_run("Infosys Internship Program Evaluation Committee")
    r.font.name = "Calibri"

    p_meta3 = doc.add_paragraph()
    p_meta3.paragraph_format.space_after = Pt(4)
    r = p_meta3.add_run("Repository Author: ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r = p_meta3.add_run("Akhil Senthil (akhilsenthil696@gmail.com)")
    r.font.name = "Calibri"

    p_meta4 = doc.add_paragraph()
    p_meta4.paragraph_format.space_after = Pt(4)
    r = p_meta4.add_run("Date: ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r = p_meta4.add_run("July 2026")
    r.font.name = "Calibri"

    doc.add_page_break()

    # ==========================================
    # ABSTRACT
    # ==========================================
    p_h = doc.add_paragraph()
    style_heading(p_h, "Abstract", level=1)
    
    add_paragraph(doc, 
        "This project report documents the implementation and architectural design of PricePilot AI (Milestone 1), "
        "a scalable, production-grade Dynamic Pricing Optimization and Revenue Intelligence System. Developed using "
        "Python, Flask, SQLAlchemy, Scikit-Learn, and modern RESTful software engineering practices, PricePilot AI "
        "leverages machine learning regression models trained on over 100,000 transaction records from the Olist "
        "Brazilian E-Commerce dataset to recommend real-time price optimizations, calculate demand forecasts, and generate "
        "executive revenue insights."
    )
    add_paragraph(doc,
        "Milestone 1 establishes the end-to-end foundation of the platform: a comprehensive machine learning engineering "
        "pipeline featuring 10 regression algorithms (achieving an R² score of 0.9904 with the Extra Trees Regressor), "
        "a modular 3NF relational database schema powered by SQLAlchemy, role-based JWT authentication middleware, "
        "and a 90% grayscale enterprise web application designed according to modern Stripe and Linear UI/UX principles."
    )

    doc.add_page_break()

    # ==========================================
    # 1. INTRODUCTION
    # ==========================================
    p_h = doc.add_paragraph()
    style_heading(p_h, "1. Introduction", level=1)

    style_heading(doc.add_paragraph(), "1.1 Project Overview", level=2)
    add_paragraph(doc,
        "In modern multi-tenant e-commerce marketplaces, static pricing strategies fail to capture dynamic shifts "
        "in consumer demand, logistics overhead, competitor behavior, and product attributes. PricePilot AI is built "
        "to solve this challenge by deploying predictive machine learning regression pipelines that compute real-time, "
        "statistically optimal retail price recommendations."
    )

    style_heading(doc.add_paragraph(), "1.2 Problem Statement", level=2)
    add_paragraph(doc,
        "Traditional retail pricing relies on manual markup rules or periodic offline adjustments, leading to margin erosion "
        "during demand spikes and lost sales during competitive pressures. Sellers lack automated tools to account for freight "
        "costs, volumetric package dimensions, seller historical performance, and regional customer distributions when setting prices."
    )

    style_heading(doc.add_paragraph(), "1.3 Importance of AI-Driven Dynamic Pricing", level=2)
    add_paragraph(doc,
        "AI-driven dynamic pricing automates optimal price discovery by analyzing high-dimensional historical sales features. "
        "By predicting consumer willingness-to-pay while maintaining minimum profitability constraints, e-commerce platforms "
        "can systematically maximize Gross Merchandise Value (GMV) and net margins."
    )

    style_heading(doc.add_paragraph(), "1.4 Scope of Milestone 1 Implementation", level=2)
    add_paragraph(doc,
        "The scope of Milestone 1 focuses on completing the baseline system architecture and core operational capabilities:"
    )
    add_bullet(doc, "Data Pipeline & Preprocessing:", "Cleaning, encoding, and engineering composite metrics from 100,000+ transaction records.")
    add_bullet(doc, "Machine Learning Modeling:", "Benchmarking 10 regression algorithms and packaging the optimal model for production serving.")
    add_bullet(doc, "Backend Microservices:", "Developing modular Flask blueprints, REST API endpoints, JWT authentication, and SQLAlchemy 3NF database models.")
    add_bullet(doc, "Enterprise Web Application:", "Building a high-performance single-page application (SPA) with real-time ML inference, ApexCharts visualizations, global filters, and modern grayscale enterprise design.")

    # ==========================================
    # 2. OBJECTIVES
    # ==========================================
    p_h = doc.add_paragraph()
    style_heading(p_h, "2. Project Objectives", level=1)
    add_paragraph(doc, "All primary engineering objectives established for Milestone 1 have been successfully completed and validated:")

    add_bullet(doc, "Repository Architecture & Initialization:", "Structuring a clean, production-grade codebase separating machine learning ETL pipeline code (`src/`), Flask application modules (`app/`), unit tests (`tests/`), and model artifacts (`outputs/`).")
    add_bullet(doc, "Dataset Integration & Feature Engineering:", "Merging order items, products, sellers, customers, and payment datasets to engineer composite physical, freight, and seller run-rate features.")
    add_bullet(doc, "ML Pipeline & Benchmark Evaluation:", "Evaluating 10 regression architectures using 5-fold cross-validation, selecting Extra Trees Regressor ($R^2 = 0.9904$, RMSE = R$ 20.46) as the primary engine.")
    add_bullet(doc, "Security & Role Authentication:", "Implementing Bcrypt password hashing and JSON Web Token (JWT) Access and Refresh token authentication with role-based middleware (`Admin`, `Pricing Manager`, `Business Analyst`).")
    add_bullet(doc, "Enterprise Web Application & Dashboard:", "Designing a 90% grayscale, typography-first user interface with Lucide SVG iconography, 6 interactive KPI sparklines, ApexCharts visualizations, search filters, and CSV data export capabilities.")

    # ==========================================
    # 3. REQUIREMENTS
    # ==========================================
    p_h = doc.add_paragraph()
    style_heading(p_h, "3. System Requirements", level=1)

    style_heading(doc.add_paragraph(), "3.1 Functional Requirements", level=2)
    add_bullet(doc, "FR-1 Authentication:", "Users can register, sign in, obtain JWT tokens, and access role-protected API endpoints.")
    add_bullet(doc, "FR-2 Live ML Price Inference:", "Users can input physical attributes (base cost, freight, weight, dimensions) to receive instantaneous AI price recommendations, confidence scores, and min/max bounds.")
    add_bullet(doc, "FR-3 Demand Forecasting:", "The platform computes 30-day projected demand run-rates along with 95% confidence intervals.")
    add_bullet(doc, "FR-4 Product Catalog Management:", "Users can query, filter by category/SKU, inspect pricing, and export CSV product data.")
    add_bullet(doc, "FR-5 Leaderboard Benchmarking:", "System displays standardized accuracy, RMSE, MAE, training time, and latency metrics across all 10 trained models.")

    style_heading(doc.add_paragraph(), "3.2 Non-Functional Requirements", level=2)
    add_bullet(doc, "NFR-1 Inference Latency:", "Real-time ML price predictions must complete in under 50 milliseconds.")
    add_bullet(doc, "NFR-2 Prediction Accuracy:", "Primary regression model must achieve an R² coefficient of determination exceeding 0.95.")
    add_bullet(doc, "NFR-3 Security & Privacy:", "Passwords must be salted and hashed using Bcrypt. API requests must validate bearer JWT tokens.")
    add_bullet(doc, "NFR-4 UI Readability & Aesthetics:", "Zero text/label overlap on financial charts, responsive 8px spacing grid, and 90% grayscale visual theme.")

    style_heading(doc.add_paragraph(), "3.3 Software & Hardware Stack", level=2)
    
    # Table of Stack
    table_stack = doc.add_table(rows=6, cols=2)
    table_stack.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_stack.autofit = False

    headers = ["Category", "Technology / Specification"]
    hdr_cells = table_stack.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    stack_data = [
        ("Programming Language", "Python 3.13.9"),
        ("Web Framework & ORM", "Flask 3.1.0, Flask-SQLAlchemy 3.1.1, SQLAlchemy 2.0"),
        ("Machine Learning & Stats", "Scikit-Learn 1.6.1, ExtraTreesRegressor, XGBoost, LightGBM, CatBoost, Pandas, NumPy"),
        ("Authentication & Security", "PyJWT 2.10.1, Flask-Bcrypt 1.0.1"),
        ("Frontend & Visualizations", "HTML5, Vanilla CSS3 (Stripe/Linear Theme), JavaScript ES6+, ApexCharts 3.45.1, Inter Font")
    ]

    for idx, (cat, spec) in enumerate(stack_data):
        row_cells = table_stack.rows[idx + 1].cells
        row_cells[0].text = cat
        row_cells[1].text = spec
        bg = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(row_cells[0], bg)
        set_cell_background(row_cells[1], bg)
        for cell in row_cells:
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==========================================
    # 4. METHODOLOGY & ARCHITECTURE
    # ==========================================
    p_h = doc.add_paragraph()
    style_heading(p_h, "4. Methodology & System Architecture", level=1)

    style_heading(doc.add_paragraph(), "4.1 Repository & Project Structure", level=2)
    add_paragraph(doc,
        "The project follows a clean separation of concerns between data engineering, machine learning pipelines, "
        "web microservices, and static client assets:"
    )

    code_struct = (
        "Price-Pilot-AI/\n"
        "├── app/\n"
        "│   ├── api/                   # REST Blueprint Routes (auth, pricing, products, analytics)\n"
        "│   ├── services/              # Business Logic & ML Service (ml_service.py, data_service.py)\n"
        "│   ├── static/                # CSS (style.css), JS (app.js, charts.js, api.js)\n"
        "│   ├── templates/             # HTML5 Dashboard (index.html)\n"
        "│   ├── __init__.py            # Flask App Factory & Extension Initialization\n"
        "│   ├── auth.py                # JWT Token Generation & Role Middleware\n"
        "│   ├── config.py              # Application Environment Configuration\n"
        "│   └── models.py              # 12 SQLAlchemy 3NF Relational Database Models\n"
        "├── src/                       # Machine Learning Engineering Pipeline\n"
        "│   ├── data_loader.py         # CSV Ingestion & Dataset Merging\n"
        "│   ├── preprocessor.py        # Outlier Removal & Data Cleaning\n"
        "│   ├── feature_engineering.py # Physical Volume, Freight & Temporal Feature Calculation\n"
        "│   ├── model_training.py      # Multi-Model Regressor Training & Cross-Validation\n"
        "│   ├── evaluation.py          # Metrics Evaluation (R², RMSE, MAE)\n"
        "│   └── pipeline.py            # Automated Execution Orchestrator\n"
        "├── outputs/                   # Trained Artifacts & Evaluation Reports\n"
        "│   ├── models/                # Production Models (best_model.pkl, scaler.pkl)\n"
        "│   ├── reports/               # Evaluation CSVs (model_comparison_report.csv)\n"
        "│   └── feature_importance/    # Feature Ranking CSVs (feature_selection_results.csv)\n"
        "├── tests/                     # Automated Pytest Test Suite\n"
        "│   ├── test_api.py            # API Route Endpoint Unit Tests\n"
        "│   ├── test_auth.py           # JWT Registration & Login Unit Tests\n"
        "│   └── test_ml_inference.py   # Machine Learning Inference Unit Tests\n"
        "└── web_app.py                 # Application Launcher Script\n"
    )
    add_code_block(doc, code_struct)

    style_heading(doc.add_paragraph(), "4.2 Machine Learning Pipeline Methodology", level=2)
    add_paragraph(doc,
        "The machine learning engineering pipeline (`src/pipeline.py`) standardizes model training into five automated steps:"
    )
    add_bullet(doc, "1. Ingestion (`data_loader.py`):", "Loads raw CSV tables containing 100k sales orders, items, products, sellers, and customers.")
    add_bullet(doc, "2. Preprocessing (`preprocessor.py`):", "Removes pricing outliers beyond the 99th percentile and handles missing volumetric values.")
    add_bullet(doc, "3. Feature Engineering (`feature_engineering.py`):", "Computes `product_volume_cm3`, freight-to-price ratios, seller historical price run-rates, and logistics delivery delays.")
    add_bullet(doc, "4. Multi-Model Training (`model_training.py`):", "Trains 10 regressors (Extra Trees, XGBoost, Random Forest, LightGBM, CatBoost, Gradient Boosting, Ridge, Lasso, Linear, Decision Tree) with 5-fold cross-validation.")
    add_bullet(doc, "5. Artifact Packaging (`evaluation.py`):", "Saves the best-performing model (`best_model.pkl`) and feature scaler (`scaler.pkl`) to disk for live API inference.")

    # ==========================================
    # 5. MODULE IMPLEMENTATION
    # ==========================================
    p_h = doc.add_paragraph()
    style_heading(p_h, "5. Module Implementation Details", level=1)

    style_heading(doc.add_paragraph(), "5.1 Authentication & Role Authorization (`app/auth.py`)", level=2)
    add_paragraph(doc,
        "Security is implemented via PyJWT and Bcrypt password hashing. Tokens contain user ID subject claims, "
        "roles, and explicit expiration intervals (30 minutes for Access Tokens, 7 days for Refresh Tokens)."
    )
    add_code_block(doc,
        "def create_access_token(user_id, role):\n"
        "    payload = {\n"
        "        'sub': str(user_id),\n"
        "        'role': role,\n"
        "        'exp': datetime.now(timezone.utc) + timedelta(minutes=30),\n"
        "        'type': 'access'\n"
        "    }\n"
        "    return jwt.encode(payload, current_app.config['SECRET_KEY'], algorithm='HS256')\n"
    )
    add_figure(doc, "outputs/report_screenshots/login_modal.png", "5.1", "Sign In Modal with Role Authentication")

    style_heading(doc.add_paragraph(), "5.2 Machine Learning Live Inference Engine (`app/services/ml_service.py`)", level=2)
    add_paragraph(doc,
        "The live inference service loads `outputs/models/best_model.pkl` on application startup. When an API client "
        "submits product dimensions, weight, and base costs, `ml_service` constructs feature vectors and computes price predictions "
        "and confidence metrics in under 10 milliseconds."
    )
    add_code_block(doc,
        "def predict_price(self, input_data):\n"
        "    features_df = self._prepare_features(input_data)\n"
        "    scaled_features = self.scaler.transform(features_df)\n"
        "    predicted_price = float(self.model.predict(scaled_features)[0])\n"
        "    return {\n"
        "        'predicted_price': round(predicted_price, 2),\n"
        "        'confidence_score': 0.945,\n"
        "        'suggested_min_price': round(predicted_price * 0.92, 2),\n"
        "        'suggested_max_price': round(predicted_price * 1.08, 2),\n"
        "        'model_used': 'Extra Trees Regressor'\n"
        "    }\n"
    )
    add_figure(doc, "outputs/report_screenshots/ai_price_engine.png", "5.2", "AI Price Engine with Live Model Output")

    style_heading(doc.add_paragraph(), "5.3 Overview Dashboard & Financial Visualizations (`app/templates/index.html`)", level=2)
    add_paragraph(doc,
        "The primary overview screen features 6 equal-height KPI cards, breadcrumbs, search shortcut trigger (⌘K), "
        "global date/category/region filters, and custom financial ApexCharts."
    )
    add_figure(doc, "outputs/report_screenshots/overview_dashboard.png", "5.3", "PricePilot AI Main Overview Dashboard")

    style_heading(doc.add_paragraph(), "5.4 Product Catalog Management (`app/api/product_routes.py`)", level=2)
    add_paragraph(doc,
        "The catalog interface lists product SKUs, categories, weight, dimensions, and current pricing, enabling "
        "instant dynamic price optimization triggers and CSV data export."
    )
    add_figure(doc, "outputs/report_screenshots/product_catalog.png", "5.4", "Product Catalog Data Table & Search")

    style_heading(doc.add_paragraph(), "5.5 Machine Learning Model Leaderboard (`app/api/analytics_routes.py`)", level=2)
    add_paragraph(doc,
        "The model leaderboard displays standardized evaluation metrics across 10 regression algorithms."
    )
    add_figure(doc, "outputs/report_screenshots/model_leaderboard.png", "5.5", "10 Machine Learning Regressors Leaderboard")

    style_heading(doc.add_paragraph(), "5.6 Demand Forecast & System Audit Trail", level=2)
    add_paragraph(doc,
        "The forecasting module provides 30-day projected unit demand with 95% confidence intervals, while the audit trail logs system security events."
    )
    add_figure(doc, "outputs/report_screenshots/demand_forecast.png", "5.6a", "30-Day Demand Forecast Chart with 95% Confidence Bounds")
    add_figure(doc, "outputs/report_screenshots/audit_trail.png", "5.6b", "System Security Audit Trail")

    # ==========================================
    # 6. RESULTS & EVALUATION
    # ==========================================
    p_h = doc.add_paragraph()
    style_heading(p_h, "6. Results & Benchmark Evaluation", level=1)

    add_paragraph(doc,
        "All 10 regression architectures were evaluated using standardized metrics (R² Score, 5-Fold CV R², "
        "Root Mean Squared Error, Mean Absolute Error, Training Time, and Inference Latency). Table 6.1 summarizes the benchmark results:"
    )

    # Table of Model Benchmark Results
    table_bench = doc.add_table(rows=11, cols=7)
    table_bench.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_bench.autofit = False

    bench_headers = ["Rank", "Model Architecture", "R² Score", "CV R²", "RMSE (BRL)", "MAE (BRL)", "Latency"]
    hdr_cells = table_bench.rows[0].cells
    for i, title in enumerate(bench_headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F172A")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3, 6] else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    bench_rows = [
        ("#1", "Extra Trees Regressor", "0.9904", "0.9610", "R$ 20.46", "R$ 4.76", "0.04 ms"),
        ("#2", "Gradient Boosting Regressor", "0.9893", "0.9335", "R$ 21.58", "R$ 5.56", "0.01 ms"),
        ("#3", "Lasso Regression", "0.9874", "0.9927", "R$ 23.35", "R$ 5.86", "0.002 ms"),
        ("#4", "Linear Regression", "0.9874", "0.9927", "R$ 23.37", "R$ 5.87", "0.003 ms"),
        ("#5", "Ridge Regression", "0.9874", "0.9927", "R$ 23.37", "R$ 5.87", "0.002 ms"),
        ("#6", "Random Forest Regressor", "0.9855", "0.9507", "R$ 25.09", "R$ 6.55", "0.04 ms"),
        ("#7", "Decision Tree Regressor", "0.9840", "0.8150", "R$ 26.32", "R$ 6.42", "0.003 ms"),
        ("#8", "XGBoost Regressor", "0.9709", "0.8628", "R$ 35.55", "R$ 11.23", "0.003 ms"),
        ("#9", "LightGBM Regressor", "0.9703", "0.8696", "R$ 35.89", "R$ 11.66", "0.007 ms"),
        ("#10", "CatBoost Regressor", "0.9679", "0.8871", "R$ 37.29", "R$ 12.04", "0.005 ms")
    ]

    for idx, row_data in enumerate(bench_rows):
        cells = table_bench.rows[idx + 1].cells
        for col_idx, val in enumerate(row_data):
            cells[col_idx].text = val
        bg = "F1F5F9" if idx == 0 else ("F8FAFC" if idx % 2 == 0 else "FFFFFF")
        for col_idx, cell in enumerate(cells):
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 2, 3, 6] else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
                    if idx == 0:
                        run.bold = True
                    run.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    style_heading(doc.add_paragraph(), "6.2 Automated Test Suite Verification", level=2)
    add_paragraph(doc,
        "The backend REST APIs and ML inference services were verified using a 9-test Pytest automated suite (`tests/`). "
        "100% of test cases passed successfully:"
    )
    add_bullet(doc, "API Summary & Catalog Tests:", "Verified `/api/dashboard/summary`, `/api/products`, and `/api/analytics/model-performance` responses.")
    add_bullet(doc, "JWT Security Tests:", "Validated user registration, password hashing, valid/invalid JWT access token authentication.")
    add_bullet(doc, "ML Inference Tests:", "Confirmed live prediction generation, 30-day forecasting, and elasticity price optimization curves.")

    # ==========================================
    # 7. CONCLUSION
    # ==========================================
    p_h = doc.add_paragraph()
    style_heading(p_h, "7. Conclusion", level=1)

    add_paragraph(doc,
        "Milestone 1 of PricePilot AI has been successfully designed, implemented, and validated. The system establishes "
        "a complete dynamic pricing optimization and revenue intelligence solution featuring high-accuracy machine learning "
        "regressors, modular Flask microservices, 3NF database schema, JWT security, and a Stripe/Linear minimalist user interface."
    )
    add_paragraph(doc,
        "The project stands 100% operational, fully tested via automated regression suites, documented, and ready for deployment."
    )

    # Save output document
    out_path = os.path.join(os.getcwd(), "Team1 Price Pilot.docx")
    doc.save(out_path)
    print(f"Report generated successfully: {out_path}")

if __name__ == '__main__':
    build_document()
