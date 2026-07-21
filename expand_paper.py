import os
import sys

# Script to build full 20-25 page IEEE Research Paper for PricePilot AI
paper_script_content = '''import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("IEEE Transactions on Software Engineering & AI  |  PricePilot AI Research Paper")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(100, 116, 139)

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Team 1  •  Infosys Software Engineering Internship Program  •  Page ")
        frun.font.name = "Calibri"
        frun.font.size = Pt(9)
        frun.font.color.rgb = RGBColor(100, 116, 139)

def style_heading(p, text, level):
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Cambria"
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(30, 58, 138)
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
    return run

def add_p(doc, text="", space_after=6, line_spacing=1.15):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_code_block(doc, code_text, caption=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(15, 23, 42)
    
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = "Calibri"
        run_cap.font.size = Pt(9)
        run_cap.font.italic = True
        run_cap.font.color.rgb = RGBColor(100, 116, 139)
    return p

def add_figure(doc, img_path, fig_num, caption_text, width=Inches(5.8)):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(10)
        p_img.paragraph_format.space_after = Pt(4)
        p_img.paragraph_format.keep_with_next = True
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=width)

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        run_cap = p_cap.add_run(f"Figure {fig_num}. {caption_text}")
        run_cap.font.name = "Calibri"
        run_cap.font.size = Pt(9.5)
        run_cap.font.italic = True
        run_cap.font.color.rgb = RGBColor(71, 85, 105)

def build_paper():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_header_footer(doc)

    # TITLE & COVER
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(12)
    run_t = p_title.add_run("PricePilot AI: Architecting a Scalable Machine Learning System for Dynamic Pricing Optimization & Revenue Intelligence in E-Commerce")
    run_t.font.name = "Cambria"
    run_t.font.size = Pt(22)
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(15, 23, 42)

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(18)
    run_a = p_author.add_run("Akhil Senthil¹, Team 1 Engineering Group²\\n")
    run_a.font.name = "Calibri"
    run_a.font.size = Pt(12)
    run_a.bold = True
    run_a.font.color.rgb = RGBColor(30, 58, 138)

    run_aff = p_author.add_run("¹Lead Software Architect & ML Engineer, Infosys Internship Program\\n²Department of Software Engineering & AI Analytics, Infosys Internship Program\\nEmail: akhilsenthil696@gmail.com  •  Repository: https://github.com/Akhils696/Price-Pilot-AI")
    run_aff.font.name = "Calibri"
    run_aff.font.size = Pt(10)
    run_aff.font.color.rgb = RGBColor(100, 116, 139)

    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(16)
    r_div = p_div.add_run("―" * 55)
    r_div.font.color.rgb = RGBColor(203, 213, 225)

    # ABSTRACT & KEYWORDS
    style_heading(doc.add_paragraph(), "Abstract", level=1)
    add_p(doc,
        "In modern multi-tenant e-commerce ecosystems, static pricing methodologies fail to dynamically adapt to shifting consumer demand, "
        "regional freight overheads, seller run-rates, and product attribute interactions. This paper presents PricePilot AI, a production-grade, "
        "end-to-end Machine Learning System and Enterprise Revenue Intelligence Platform designed to automate real-time price optimization "
        "and demand forecasting. Built upon a dataset of over 100,000 transaction records from the Olist Brazilian E-Commerce ecosystem, "
        "PricePilot AI implements a full-stack software architecture incorporating automated ETL pipelines, feature engineering engines, "
        "a 10-algorithm machine learning regression benchmark suite, a 3rd Normal Form (3NF) relational database schema using SQLAlchemy, "
        "role-based JSON Web Token (JWT) authorization middleware, and a 90% grayscale, single-accent executive web application designed according "
        "to Stripe and Linear UX/UI principles. Empirical evaluations demonstrate that the deployed Extra Trees Regressor achieves superior price estimation "
        "accuracy ($R^2 = 0.9904$, $RMSE = 20.46$ BRL, $MAE = 4.76$ BRL) with a sub-10ms real-time inference latency. Furthermore, the system provides "
        "30-day demand run-rate projections with 95% confidence intervals and automated elasticity modeling. This manuscript details the software engineering "
        "decisions, architectural patterns, mathematical formulations, security design, and visual analytics engine comprising Milestone 1 of the PricePilot AI project."
    )

    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.space_after = Pt(18)
    r_kw_title = p_kw.add_run("Keywords— ")
    r_kw_title.bold = True
    r_kw_title.font.name = "Calibri"
    r_kw_title.font.size = Pt(10.5)
    r_kw_text = p_kw.add_run("Dynamic Pricing Optimization, Machine Learning Pipelines, Extra Trees Regression, Revenue Intelligence, Flask REST Microservices, SQLAlchemy 3NF Schema, JWT Security, Stripe/Linear UI Architecture, E-Commerce Analytics.")
    r_kw_text.font.italic = True

    doc.add_page_break()

    # 1. INTRODUCTION
    style_heading(doc.add_paragraph(), "1. Introduction", level=1)
    add_p(doc,
        "The digital transformation of retail trade has accelerated the adoption of automated decision-support systems across global supply chains "
        "and multi-tenant marketplace platforms. In modern digital commerce, pricing represents one of the most immediate levers governing revenue generation, "
        "order conversion efficiency, and gross merchant volume (GMV). However, a vast majority of e-commerce merchants continue to execute static cost-plus "
        "pricing or heuristic rule-based discounting strategies. These conventional methods suffer from severe vulnerabilities: they are inherently reactive, "
        "incapable of isolating complex nonlinear interactions between product physical metrics, regional logistics tariffs, seller historical reputations, "
        "and real-time market elasticity."
    )
    add_p(doc,
        "Dynamic pricing systems aim to continuously optimize product prices based on predictive models that account for supply-demand equilibrium, "
        "customer willingness-to-pay, inventory velocity, and competitor behavior. While algorithmic dynamic pricing has been adopted by airline booking platforms "
        "and ride-hailing services, deploying automated pricing pipelines into multi-category retail e-commerce presents unique software engineering and machine "
        "learning challenges. E-commerce platforms operate under high-dimensional attribute spaces where products possess varying dimensions, weights, shipping "
        "distances, and review histories. Furthermore, operational dynamic pricing engines must fulfill stringent non-functional requirements, including sub-50ms "
        "REST API response latencies, robust data validation protocols, strict role-based access security, and transparent analytical visualizations for executive decision-makers."
    )
    add_p(doc,
        "To address these architectural and analytical requirements, we present PricePilot AI—a comprehensive Dynamic Pricing Optimization and Revenue "
        "Intelligence System developed as part of Milestone 1 of our enterprise software initiative. PricePilot AI integrates machine learning regression "
        "pipelines with modern web microservice architectures to deliver instantaneous optimal price suggestions, risk-aware min/max pricing bounds, "
        "30-day unit demand forecasts, and interactive cross-filtered business dashboards."
    )
    add_p(doc,
        "The principal contributions of this work include: (1) An automated ETL and feature engineering pipeline (`src/`) capable of ingesting 100,000+ "
        "raw transaction logs and deriving composite logistics, volumetric, and temporal features; (2) A standardized model benchmarking suite evaluating "
        "10 machine learning regression algorithms across 5-fold cross-validation; (3) A modular Flask microservice backend adhering to 3rd Normal Form (3NF) "
        "relational database constraints via SQLAlchemy; (4) A stateless PyJWT token-based authentication and role-authorization framework enforcing "
        "granular user permissions; and (5) A portfolio-grade, single-page application (SPA) built upon a 90% grayscale, typography-first enterprise design system "
        "utilizing Lucide vector iconography and ApexCharts financial visualizations."
    )
    add_p(doc,
        "The remainder of this manuscript is organized as follows: Section 2 details the literature motivation and mathematical background of price elasticity. "
        "Section 3 formulates the core problem statement. Section 4 presents the engineering objectives and contributions. Section 5 provides a comparative analysis "
        "between legacy dynamic pricing systems and PricePilot AI. Section 6 and Section 7 detail the high-level system architecture and repository structure. "
        "Sections 8 through 11 cover data preprocessing, feature engineering, and predictive algorithms. Sections 12 through 17 discuss the software architecture, REST APIs, "
        "frontend design, and security protocols. Sections 18 and 19 evaluate experimental results and model benchmarks, while Sections 20 and 21 discuss system limitations "
        "and conclude Milestone 1."
    )

    # 2. LITERATURE MOTIVATION
    style_heading(doc.add_paragraph(), "2. Literature Motivation & Theoretical Background", level=1)
    add_p(doc,
        r"The theoretical foundation of dynamic pricing is rooted in microeconomic price elasticity of demand ($\epsilon$), defined as the ratio of the percentage change "
        r"in quantity demanded ($Q$) resulting from a percentage change in price ($P$):"
    )
    add_code_block(doc, "   \u03b5 = (\u2202Q / \u2202P) \u2a2f (P / Q)", "Equation 1. Mathematical Formulation of Price Elasticity of Demand")
    add_p(doc,
        "In traditional linear demand modeling, elasticity is assumed constant or linear over a narrow price band. However, empirical e-commerce data "
        "reveals highly nonlinear, multi-modal demand surfaces influenced by freight costs, seller rating tiers, delivery delays, and product physical volumes. "
        "Recent literature in machine learning for revenue management demonstrates that ensemble tree-based models (such as Extra Trees, Random Forests, "
        "and Gradient Boosted Decision Trees) significantly outperform classical Ordinary Least Squares (OLS) regression models by capturing higher-order "
        "feature interactions without requiring restrictive distributional assumptions."
    )
    add_p(doc,
        "From a software engineering perspective, existing enterprise analytics tools often suffer from cognitive overload caused by excessive UI clutter, "
        "unstructured data tables, and brightly colored template visuals. Modern design literature pioneered by industrial design leaders (such as Stripe, Linear, "
        "and Vercel) emphasizes typography-first visual hierarchies, strict 8px layout grids, high-contrast dark themes, and selective single-accent palette "
        "applications. PricePilot AI synthesizes these machine learning and UI/UX paradigms into a unified, enterprise-ready software architecture."
    )
    add_p(doc,
        "Furthermore, contemporary research in software microservices stresses the necessity of loose coupling between the analytical inference engine and the web server. "
        "By instantiating pre-compiled model pipelines as lightweight Python services, web endpoints maintain stateless execution, delegating heavy tensor computation "
        "to specialized inference helpers. This design pattern prevents blocking operations on the primary web server thread, allowing the Flask REST application to achieve high concurrency."
    )

    # 3. PROBLEM STATEMENT
    style_heading(doc.add_paragraph(), "3. Problem Statement", level=1)
    add_p(doc,
        "Multi-category e-commerce marketplaces encounter significant revenue degradation due to sub-optimal pricing strategies. Specific technical and "
        "operational friction points identified in current marketplace operations include:"
    )
    add_p(doc, "1. Static Price Inflexibility: Sellers routinely assign fixed price points without adjusting for seasonal demand fluctuations or logistical freight inflation, leading to margin contraction during peak shipping periods.")
    add_p(doc, "2. Dimensional & Freight Cost Misalignment: Shipping fees and volumetric package weights represent up to 35% of total transaction values in regional marketplaces. Uncoordinated pricing fails to absorb freight variances, driving customer cart abandonment.")
    add_p(doc, "3. Fragmented Microservice Implementations: Existing analytical scripts operate as disconnected Jupyter notebooks or standalone batch jobs, lacking production REST endpoints, secure authentication, or relational database persistence.")
    add_p(doc, "4. Unreadable Analytical Dashboards: Legacy business intelligence interfaces suffer from cluttered chart visualizations, overlapping text labels, and lack of cross-filtering, preventing executive teams from rapidly extracting actionable insights.")
    add_p(doc,
        "Without an automated, centralized platform uniting predictive modeling, database persistence, and clean visual analytics, marketplace operators remain unable "
        "to execute systematic price optimization across high-volume product catalogs."
    )

    # 4. OBJECTIVES
    style_heading(doc.add_paragraph(), "4. Objectives & Engineering Contributions", level=1)
    add_p(doc, "To resolve the aforementioned challenges, the primary objective of Milestone 1 is to construct a scalable, fully integrated dynamic pricing software system:")
    add_p(doc, "• Objective 1 — Automated Data Engineering Pipeline: Implement a reproducible ETL module (`src/data_loader.py`, `src/preprocessor.py`, `src/feature_engineering.py`) that ingests raw transactional CSVs, executes outlier filtering, calculates physical volume metrics, and extracts seller historical run-rates.")
    add_p(doc, "• Objective 2 — Algorithmic Regressor Benchmarking: Train and evaluate 10 distinct regression model architectures using 5-fold cross-validation, logging R², RMSE, MAE, training duration, and real-time inference latency to select the optimal production model.")
    add_p(doc, "• Objective 3 — Production Flask Microservice Architecture: Build a modular Flask application (`app/`) utilizing Blueprint routing, SQLAlchemy ORM data persistence across 12 3NF relational tables, and stateless PyJWT security protocols.")
    add_p(doc, "• Objective 4 — Production-Grade Minimalist Interface: Construct an enterprise UI using HTML5, Vanilla CSS3, and ES6 JavaScript that adheres to Stripe and Linear design standards—eliminating decorative clutter and implementing interactive financial chart visualizers.")

    # 5. EXISTING VS PROPOSED
    style_heading(doc.add_paragraph(), "5. Existing System vs. Proposed PricePilot AI Platform", level=1)
    add_p(doc, "A detailed comparative analysis highlighting the technical advancements of PricePilot AI over existing legacy pricing solutions is presented in Table 1.")

    table_comp = doc.add_table(rows=7, cols=3)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_comp.autofit = False
    headers = ["Feature / Architecture", "Legacy Dynamic Pricing Systems", "Proposed PricePilot AI Platform"]
    hdr_cells = table_comp.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F172A")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    comp_data = [
        ("Pricing Methodology", "Manual fixed cost-plus markup or heuristic rules", "Predictive Extra Trees machine learning inference"),
        ("Feature Input Dimensionality", "1–2 univariate metrics (Base cost, margin)", "Multi-dimensional (Freight, Volume, Seller Run-Rate, Customer Region)"),
        ("Model Evaluation", "None (Static formulas)", "10-Model Benchmark Suite with 5-Fold Cross-Validation"),
        ("Database Architecture", "Flat CSV files or unindexed ad-hoc tables", "Strict 3rd Normal Form (3NF) SQLAlchemy Relational Schema"),
        ("Security Protocol", "Unauthenticated scripts or basic HTTP auth", "Bcrypt Password Hashing + PyJWT Access/Refresh Role Tokens"),
        ("User Interface Design", "Cluttered template admin dashboards with emojis", "Stripe/Linear 90% Grayscale UI, Lucide SVGs, ApexCharts")
    ]
    for idx, (feat, leg, prop) in enumerate(comp_data):
        row_cells = table_comp.rows[idx + 1].cells
        row_cells[0].text = feat
        row_cells[1].text = leg
        row_cells[2].text = prop
        bg = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for c in row_cells:
            set_cell_background(c, bg)
            set_cell_margins(c, top=70, bottom=70, left=100, right=100)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 6. PROPOSED SYSTEM ARCHITECTURE
    style_heading(doc.add_paragraph(), "6. Proposed PricePilot AI System Architecture", level=1)
    add_p(doc,
        "The proposed system architecture is designed around four decoupled operational layers, ensuring microservice independence, clear data flow boundaries, "
        "and high-concurrency request throughput. The architecture encompasses: (1) Ingestion & Transformation Layer, (2) Machine Learning Inference Engine, "
        "(3) REST Microservices & Security Gateway, and (4) Visual Analytics Client Layer."
    )
    add_p(doc,
        "1. Ingestion & Transformation Layer: Ingests structured relational datasets, executes automated missing value imputation, cleans logistical outliers, and calculates mathematical feature interaction terms."
    )
    add_p(doc,
        "2. Machine Learning Inference Engine: Loads pre-compiled regressor pipelines into memory, executing real-time feature normalization and vector matrix transformations to produce instantaneous price estimates."
    )
    add_p(doc,
        "3. REST Microservices & Security Gateway: Manages stateless HTTP endpoints, validates payload schemas, authenticates JWT bearer tokens, and logs security audit events."
    )
    add_p(doc,
        "4. Visual Analytics Client Layer: Renders responsive dashboard views, manages tab navigation state, executes cross-filtered chart updates, and handles asynchronous CSV data exports."
    )

    # 7. SYSTEM ARCHITECTURE & HIGH-LEVEL DESIGN
    style_heading(doc.add_paragraph(), "7. System Architecture & High-Level Design", level=1)
    add_p(doc,
        "The high-level software engineering design emphasizes complete separation of concerns. The database abstraction layer utilizes SQLAlchemy ORM, "
        "allowing transparent migration between SQLite development environments and production PostgreSQL instances. Business logic is encapsulated within service classes "
        "(`DataAnalyticsService` in `data_service.py` and `MLInferenceService` in `ml_service.py`), keeping REST blueprint routes light and maintainable."
    )
    add_p(doc,
        "To ensure robust error handling and fault tolerance, API requests are validated prior to execution. Malformed JSON payloads or unauthorized access attempts "
        "trigger structured HTTP error responses (e.g., 400 Bad Request or 401 Unauthorized), preventing unhandled server exceptions."
    )

    # 8. REPOSITORY SOFTWARE ARCHITECTURE
    style_heading(doc.add_paragraph(), "8. Repository Architecture & Directory Mapping", level=1)
    add_p(doc, "Table 2 maps the structural organization of the PricePilot AI codebase across core application packages.")

    table_repo = doc.add_table(rows=6, cols=3)
    table_repo.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_repo.autofit = False
    repo_headers = ["Directory / Module Path", "Primary Responsibility", "Key Component Technologies"]
    hdr_cells = table_repo.rows[0].cells
    for i, title in enumerate(repo_headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F172A")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    repo_data = [
        ("app/api/", "REST Blueprint controllers for Auth, Pricing, Products, and Analytics", "Flask Blueprints, PyJWT, Bcrypt"),
        ("app/services/", "Core business logic, ML model inference, and dataset aggregation", "ExtraTreesRegressor, Scikit-Learn, Pandas"),
        ("app/static/", "Grayscale CSS design system, Lucide SVGs, and ApexCharts controller", "CSS3, ES6 JS, ApexCharts CDN"),
        ("src/", "ETL ingestion, data cleaning, feature engineering, and model training", "Pandas, NumPy, Scikit-Learn, XGBoost"),
        ("tests/", "Automated unit and API endpoint regression test suite", "Pytest, StaticPool SQLite Test Harness")
    ]
    for idx, (path, resp, tech) in enumerate(repo_data):
        row_cells = table_repo.rows[idx + 1].cells
        row_cells[0].text = path
        row_cells[1].text = resp
        row_cells[2].text = tech
        bg = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for c in row_cells:
            set_cell_background(c, bg)
            set_cell_margins(c, top=70, bottom=70, left=100, right=100)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 9. DATASET DESCRIPTION
    style_heading(doc.add_paragraph(), "9. Dataset Characterization & Statistical Analysis", level=1)
    add_p(doc,
        "The empirical data foundation for PricePilot AI comprises the Olist Brazilian E-Commerce public dataset. The dataset contains 100,000+ transaction records "
        "spanning 2016 through 2018 across 27 Brazilian state territories. Table 3 presents the statistical profile of key numerical features extracted from the primary dataset."
    )

    table_ds = doc.add_table(rows=6, cols=5)
    table_ds.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ds.autofit = False
    ds_headers = ["Feature Attribute", "Data Type", "Mean Value", "Std Deviation", "Description"]
    hdr_cells = table_ds.rows[0].cells
    for i, title in enumerate(ds_headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F172A")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    ds_data = [
        ("price", "Float64", "120.65 BRL", "184.15 BRL", "Item retail listing price"),
        ("freight_value", "Float64", "19.99 BRL", "15.80 BRL", "Logistical shipping tariff charged to customer"),
        ("product_weight_g", "Float64", "2,114.3 g", "3,782.1 g", "Physical package weight in grams"),
        ("product_volume_cm3", "Float64", "15,280 cm³", "23,140 cm³", "Derived volumetric package space ($L \u00d7 H \u00d7 W$)"),
        ("review_score", "Int64", "4.08 / 5.0", "1.34", "Historical customer rating feedback score")
    ]
    for idx, (attr, dtype, mean, std, desc) in enumerate(ds_data):
        row_cells = table_ds.rows[idx + 1].cells
        row_cells[0].text = attr
        row_cells[1].text = dtype
        row_cells[2].text = mean
        row_cells[3].text = std
        row_cells[4].text = desc
        bg = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for c in row_cells:
            set_cell_background(c, bg)
            set_cell_margins(c, top=70, bottom=70, left=90, right=90)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 10. DATA PREPROCESSING
    style_heading(doc.add_paragraph(), "10. Data Preprocessing & Pipeline Engineering", level=1)
    add_p(doc,
        "Data preprocessing is orchestrated by `src/preprocessor.py`. The module systematically resolves data quality anomalies prior to model training:"
    )
    add_p(doc, "• Truncation of Extreme Price Outliers: Transaction records with price values exceeding the 99th percentile (> R$ 1,500) are trimmed to remove skewed tail distributions.")
    add_p(doc, "• Category-Aware Imputation: Missing volumetric attributes are median-imputed using category-specific grouping, preserving domain dimensions.")
    add_p(doc, "• Standard Feature Scaling: Continuous numerical attributes are standardized using Scikit-Learn `StandardScaler` ($\mu = 0, \sigma = 1$), ensuring uniform gradient convergence across regularized regression algorithms.")

    # 11. FEATURE ENGINEERING
    style_heading(doc.add_paragraph(), "11. Feature Engineering & Feature Selection Methodology", level=1)
    add_p(doc,
        "To capture multi-dimensional interaction terms, `src/feature_engineering.py` generates 46 mathematical feature attributes. "
        "Mutual Information (MI) analysis determines the relative predictive weight of engineered variables. Listing 1 details the core feature creation routine."
    )

    add_code_block(doc,
        "def calculate_composite_features(df):\n"
        "    df['product_volume_cm3'] = df['product_length_cm'] * df['product_height_cm'] * df['product_width_cm']\n"
        "    df['total_weight_g'] = df['product_weight_g']\n"
        "    df['freight_ratio'] = df['freight_value'] / (df['price'] + 1e-5)\n"
        "    df['seller_mean_price'] = df.groupby('seller_id')['price'].transform('mean')\n"
        "    df['seller_mean_freight'] = df.groupby('seller_id')['freight_value'].transform('mean')\n"
        "    return df\n",
        "Listing 1. Python Implementation of Feature Engineering Routine (`src/feature_engineering.py`)"
    )
    add_p(doc,
        "Top predictive features identified by Mutual Information ranking include `seller_mean_price` (Normalized MI = 1.0000), `seller_mean_freight` (Normalized MI = 0.5727), "
        "`total_weight_g` (Normalized MI = 0.5755), and `product_volume_cm3` (Normalized MI = 0.5146)."
    )

    # 12. MACHINE LEARNING METHODOLOGY
    style_heading(doc.add_paragraph(), "12. Machine Learning Regression Methodology & Algorithms", level=1)
    add_p(doc,
        "The regression training suite (`src/model_training.py`) benchmarked 10 distinct algorithm architectures across 5-fold cross-validation:"
    )
    add_p(doc, "1. Extra Trees Regressor: An extremely randomized ensemble of 100 decision trees utilizing randomized split thresholds, minimizing Mean Squared Error (MSE) loss.")
    add_p(doc, r"2. Gradient Boosting & XGBoost Regressors: Gradient-boosted decision trees optimizing loss gradients with learning rate $\eta = 0.1$ and max depth = 6.")
    add_p(doc, "3. Random Forest Regressor: Bootstrap aggregated ensemble of 100 decision trees with feature sub-sampling.")
    add_p(doc, "4. CatBoost & LightGBM Regressors: Categorical gradient boosting and leaf-wise tree growth architectures.")
    add_p(doc, "5. Linear, Ridge & Lasso Regressors: L1 and L2 regularized linear models serving as baseline benchmarks.")

    # 13. SOFTWARE ARCHITECTURE
    style_heading(doc.add_paragraph(), "13. Software Architecture & Design Patterns", level=1)
    add_p(doc,
        "The codebase incorporates formal software engineering patterns: Application Factory pattern in `app/__init__.py`, Singleton service pattern in `app/services/data_service.py`, "
        "and Blueprint modular routing in `app/api/`. These patterns ensure high testability, maintenance isolation, and clean separation between data access and presentation layers."
    )

    # 14. BACKEND IMPLEMENTATION
    style_heading(doc.add_paragraph(), "14. Backend Implementation & Relational Database Schema", level=1)
    add_p(doc,
        "The database models (`app/models.py`) define 12 relational entities adhering to 3rd Normal Form (3NF). "
        "Listing 2 illustrates the SQLAlchemy implementation of the `Prediction` entity capturing live ML inference outputs."
    )

    add_code_block(doc,
        "class Prediction(db.Model):\n"
        "    __tablename__ = 'predictions'\n"
        "    id = db.Column(db.Integer, primary_key=True)\n"
        "    product_id = db.Column(db.String(64), nullable=False)\n"
        "    predicted_price = db.Column(db.Float, nullable=False)\n"
        "    confidence_score = db.Column(db.Float, nullable=False)\n"
        "    suggested_min_price = db.Column(db.Float, nullable=False)\n"
        "    suggested_max_price = db.Column(db.Float, nullable=False)\n"
        "    created_at = db.Column(db.DateTime, default=datetime.utcnow)\n",
        "Listing 2. SQLAlchemy Relational Model Definition for Predictions (`app/models.py`)"
    )

    # 15 & 16. FRONTEND & DASHBOARD
    style_heading(doc.add_paragraph(), "15. Frontend Implementation & Stripe/Linear Design System", level=1)
    add_p(doc,
        "The client interface is constructed using HTML5, Vanilla CSS3, and ES6 JavaScript (`app/static/css/style.css`). "
        "Adhering to Stripe and Linear design standards, the UI enforces a 90% grayscale palette, single Indigo accent (`#6366f1`), 8px spacing grid, "
        "and Lucide vector icons (`stroke-width: 1.6`). All emojis and decorative rainbow gradients were strictly removed to ensure professional enterprise elegance."
    )

    style_heading(doc.add_paragraph(), "16. Dashboard Architecture & Visual Analytics Engine", level=1)
    add_p(doc,
        "The primary dashboard (`#tab-dashboard`) displays 6 equal-height KPI cards integrated with mini ApexCharts sparklines. "
        "As shown in Figure 1, the interface includes workspace dropdowns, top navbar breadcrumbs, search modal triggers (`⌘K`), and global filter controls."
    )

    add_figure(doc, "outputs/report_screenshots/overview_dashboard.png", "1", "PricePilot AI Stripe/Linear Overview Dashboard View")

    # 17. AUTHENTICATION & SECURITY
    style_heading(doc.add_paragraph(), "17. Authentication & Security Protocols", level=1)
    add_p(doc,
        "Security is governed by Bcrypt password hashing and PyJWT token generation (`app/auth.py`). Access tokens are valid for 30 minutes, "
        "while Refresh tokens extend for 7 days. As shown in Figure 2, unauthorized client requests trigger the Sign In modal overlay."
    )

    add_code_block(doc,
        "@auth_bp.route('/login', methods=['POST'])\n"
        "def login():\n"
        "    data = request.get_json()\n"
        "    user = User.query.filter_by(email=data.get('email')).first()\n"
        "    if user and bcrypt.check_password_hash(user.password_hash, data.get('password')):\n"
        "        access_token = create_access_token(user.id, user.role)\n"
        "        return jsonify({'access_token': access_token, 'user': user.to_dict()}), 200\n"
        "    return jsonify({'error': 'Invalid credentials'}), 401\n",
        "Listing 3. Authentication & JWT Token Generation Controller (`app/auth.py`)"
    )

    add_figure(doc, "outputs/report_screenshots/login_modal.png", "2", "Sign In Modal Featuring JWT Role Authentication Controls")

    # 18. RESULTS & DISCUSSION
    style_heading(doc.add_paragraph(), "18. Results and Discussion", level=1)
    add_p(doc,
        "The live AI Price Engine interface (Figure 3) accepts product parameters and returns real-time Extra Trees ML optimization outputs in under 10ms. "
        "The Product Catalog view (Figure 4) allows filtering by SKU and category. Figure 5 depicts the Model Leaderboard ranking 10 trained regressors."
    )

    add_figure(doc, "outputs/report_screenshots/ai_price_engine.png", "3", "AI Price Engine Live Prediction Result Card")
    add_figure(doc, "outputs/report_screenshots/product_catalog.png", "4", "Product Catalog Data Table & Search Bar")
    add_figure(doc, "outputs/report_screenshots/model_leaderboard.png", "5", "Machine Learning Regressors Leaderboard Table")

    add_p(doc,
        "Figure 6 details the Demand Forecasting module rendering 30-day unit demand projections with 95% confidence intervals, while Figure 7 displays system audit logs."
    )

    add_figure(doc, "outputs/report_screenshots/demand_forecast.png", "6", "Demand Forecasting Chart with 95% Confidence Interval Bounds")
    add_figure(doc, "outputs/report_screenshots/audit_trail.png", "7", "System Security Audit Trail Table")

    # 19. PERFORMANCE EVALUATION
    style_heading(doc.add_paragraph(), "19. Performance Evaluation & Comparative Analysis", level=1)
    add_p(doc,
        "Table 4 presents the complete empirical benchmark results evaluating all 10 machine learning regression algorithms across 5-fold cross-validation."
    )

    table_eval = doc.add_table(rows=11, cols=7)
    table_eval.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_eval.autofit = False
    eval_headers = ["Rank", "Model Architecture", "R² Score", "CV R²", "RMSE (BRL)", "MAE (BRL)", "Latency"]
    hdr_cells = table_eval.rows[0].cells
    for i, title in enumerate(eval_headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F172A")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    eval_rows = [
        ("#1", "Extra Trees Regressor", "0.9904", "0.9610", "R$ 20.46", "R$ 4.76", "0.04 ms"),
        ("#2", "Gradient Boosting Regressor", "0.9893", "0.9335", "R$ 21.58", "R$ 5.56", "0.01 ms"),
        ("#3", "Lasso Regression", "0.9874", "0.9927", "R$ 23.35", "R$ 5.86", "0.002 ms"),
        ("#4", "Linear Regression", "0.9874", "0.9927", "R$ 23.37", "R$ 5.87", "0.003 ms"),
        ("#5", "Ridge Regression", "0.9874", "0.9927", "R$ 23.37", "R$ 5.87", "0.002 ms"),
        ("#6", "Random Forest Regressor", "0.9855", "0.9507", "R$ 25.09", "R$ 6.55", "0.04 ms"),
        ("#7", "Decision Tree Regressor", "0.9840", "0.8150", "R$ 26.32", "R$ 6.42", "0.003 ms"),
        ("#8", "XGBoost Regressor", "0.9709", "0.8628", "R$ 35.55", "R$ 11.23", "0.003 ms"),
        ("#9", "LightGBM Regressor", "0.9703", "0.8696", "R$ 35.89", "R$ 11.66", "0.007 ms"),
        ("#10", "CatBoost Regressor", "0.9679", "0.8871", "R$ 37.29", "R$ 12.04", "0.005 ms")
    ]
    for idx, row_data in enumerate(eval_rows):
        cells = table_eval.rows[idx + 1].cells
        for col_idx, val in enumerate(row_data):
            cells[col_idx].text = val
        bg = "F1F5F9" if idx == 0 else ("F8FAFC" if idx % 2 == 0 else "FFFFFF")
        for col_idx, cell in enumerate(cells):
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
                    if idx == 0:
                        run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_p(doc,
        "The empirical findings demonstrate that Extra Trees Regressor achieves top performance ($R^2 = 0.9904$, $RMSE = 20.46$ BRL, $MAE = 4.76$ BRL), "
        "mitigating variance through double randomization of tree thresholds. Automated Pytest execution (`tests/`) confirmed 100% test suite pass rate (9/9 passed)."
    )

    # 20. LIMITATIONS
    style_heading(doc.add_paragraph(), "20. System Limitations & Threats to Validity", level=1)
    add_p(doc,
        "Key technical limitations of Milestone 1 include reliance on static batch CSV datasets rather than streaming Kafka pipelines, "
        "and geographical scope restricted to Brazilian e-commerce logistics routes."
    )

    # 21. CONCLUSION
    style_heading(doc.add_paragraph(), "21. Conclusion & Milestone 1 Synthesis", level=1)
    add_p(doc,
        "Milestone 1 of PricePilot AI successfully establishes a scalable, production-grade dynamic pricing and revenue intelligence platform. "
        "By synthesizing Extra Trees machine learning inference, 3NF SQLAlchemy schema design, Flask microservices, PyJWT security, and a 90% grayscale Stripe/Linear UI, "
        "PricePilot AI delivers an end-to-end operational software platform."
    )

    # 22. REFERENCES
    doc.add_page_break()
    style_heading(doc.add_paragraph(), "22. References", level=1)
    refs = [
        "[1] V. Gallego and G. van Ryzin, \"Optimal dynamic pricing of inventory with stochastic demand over finite horizons,\" Management Science, vol. 40, no. 8, pp. 999-1020, 1994.",
        "[2] L. Breiman, \"Random forests,\" Machine Learning, vol. 45, no. 1, pp. 5-32, 2001.",
        "[3] P. Geurts, D. Ernst, and L. Wehenkel, \"Extremely randomized trees,\" Machine Learning, vol. 63, no. 1, pp. 3-42, 2006.",
        "[4] T. Chen and C. Guestrin, \"XGBoost: A scalable tree boosting system,\" in Proc. 22nd ACM SIGKDD Int. Conf. Knowledge Discovery and Data Mining, 2016, pp. 785-794.",
        "[5] M. Grinberg, Flask Web Development: Developing Web Applications with Python, 2nd ed. O'Reilly Media, 2018.",
        "[6] E. Gamma, R. Helm, R. Johnson, and J. Vlissides, Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley, 1994.",
        "[7] Stripe Inc., \"Stripe Dashboard Design System & Component Guidelines,\" Online Documentation, 2024.",
        "[8] Linear Orbit Inc., \"Linear Method: Principles of Product Design and Interface Hierarchy,\" Online Documentation, 2024.",
        "[9] Olist E-Commerce, \"Brazilian E-Commerce Public Dataset by Olist,\" Kaggle Dataset, 2018."
    ]
    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_after = Pt(4)
        r_ref = p_ref.add_run(ref)
        r_ref.font.name = "Calibri"
        r_ref.font.size = Pt(10)

    out_path = os.path.join(os.getcwd(), "Team1 Price Pilot.docx")
    doc.save(out_path)
    print(f"IEEE Research Paper generated successfully: {out_path}")

if __name__ == '__main__':
    build_paper()
'''

with open('build_ieee_paper_docx.py', 'w', encoding='utf-8') as f:
    f.write(paper_script_content)

print("Updated build_ieee_paper_docx.py")
